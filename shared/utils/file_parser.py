"""
文件文本提取工具 - 支持 .txt, .md, .docx, .pdf

Usage:
    from shared.utils.file_parser import extract_text_from_file

    text = await extract_text_from_file(file_bytes, "interview_questions.pdf")
"""
import io
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf"}
MAX_FILE_SIZE_MB = 50


def get_file_extension(filename: str) -> str:
    """获取文件扩展名（小写）"""
    return Path(filename).suffix.lower()


def is_supported_file(filename: str) -> bool:
    """检查文件类型是否支持"""
    return get_file_extension(filename) in SUPPORTED_EXTENSIONS


async def extract_text_from_file(
    file_content: bytes,
    filename: str,
    encoding: str = "utf-8",
) -> str:
    """
    从文件中提取纯文本内容。

    Args:
        file_content: 文件二进制内容
        filename: 文件名（用于判断类型）
        encoding: 文本文件编码（默认 utf-8）

    Returns:
        提取出的纯文本

    Raises:
        ValueError: 不支持的文件类型
        RuntimeError: 文本提取失败
    """
    ext = get_file_extension(filename)

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"不支持的文件类型: {ext}，支持的类型: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    file_size_mb = len(file_content) / (1024 * 1024)
    if file_size_mb > MAX_FILE_SIZE_MB:
        raise ValueError(f"文件大小 {file_size_mb:.1f}MB 超过限制 {MAX_FILE_SIZE_MB}MB")

    try:
        if ext in (".txt", ".md"):
            return _extract_text_plain(file_content, encoding)
        elif ext == ".docx":
            return _extract_text_docx(file_content)
        elif ext == ".pdf":
            return _extract_text_pdf(file_content)
    except ValueError:
        raise
    except Exception as e:
        logger.error(f"文本提取失败 ({filename}): {e}")
        raise RuntimeError(f"文件文本提取失败: {e}")


def _extract_text_plain(content: bytes, encoding: str = "utf-8") -> str:
    """提取纯文本 / Markdown 文件"""
    try:
        return content.decode(encoding)
    except UnicodeDecodeError:
        for fallback in ("gbk", "gb2312", "latin-1"):
            try:
                return content.decode(fallback)
            except UnicodeDecodeError:
                continue
        raise RuntimeError("无法识别文件编码")


def _extract_text_docx(content: bytes) -> str:
    """提取 .docx 文件文本"""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("缺少 python-docx 依赖，请安装: pip install python-docx")

    doc = Document(io.BytesIO(content))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def _extract_text_pdf(content: bytes) -> str:
    """提取 .pdf 文件文本"""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise RuntimeError("缺少 pypdf 依赖，请安装: pip install pypdf")

    reader = PdfReader(io.BytesIO(content))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append(text.strip())

    if not pages:
        raise RuntimeError("PDF 文件未提取到文本内容，可能是扫描件或图片 PDF")

    return "\n\n".join(pages)
